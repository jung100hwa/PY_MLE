from docx import Document
from docx.shared import Inches, Cm

dictxt = {'[H01]':'김철수', '[H02]':'chemp','[H03]':'2022.12.30'}

records = (
    ('1', 'AA', '101', 'Spam','NA'),
    ('2', 'BB', '422', 'Eggs','NA'),
    ('3', 'CC', '631', 'Spam, spam, eggs, and spam','NA')
)

doc = Document("시험성적서 원본.docx")
par = doc.paragraphs

# 문서안에 값 적기
for pt in par:
    if pt.text:
        for item in dictxt.keys():
            pt.text = pt.text.replace(item,dictxt[item])

# for pt in par:
#     for item in dictxt.keys():
#         orgtext = dictxt[item]
#         if item in pt.text:
#             inline = pt.runs
#             for i in range(len(inline)):
#                 if orgtext in inline[i].text:
#                     text = inline[i].text.replace(item, orgtext)
#                     inline[i].text = text

# 문서에 포함된 테이블 리스트 불러오기
tables = doc.tables
    
# 테이블에 값넣기
table = tables[0]
rowCount = 0
row_cell = table.rows[1].cells

for item in records:
    if rowCount > 0:
        row_cell = table.add_row().cells
        for i in range(0,len(item)):
            row_cell[i].text = item[i]
    else:
        for i in range(0,len(item)):
            row_cell[i].text = item[i]
            
    rowCount = rowCount + 1

# 서명파일 넣기
cell = tables[1].rows[0].cells[2]
paragraph = cell.paragraphs[0]
run = paragraph.add_run()
run.add_picture('서명파일.jpg', width= Cm(2), height= Cm(1))

doc.save("시험성적서.docx")