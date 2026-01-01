#!/usr/bin/env python 3.9
# -*- coding: utf-8 -*-
# @Time    : 2022/12/6 18:02 update
# @Author  : ZCG
# @File    : WordReplace.py
# @Software: PyCharm
# @Notice  :
from docx import Document
import os
from docx.shared import Inches, Cm

# 실제 클래스 WordReplace의 실행 클래스
class Execute:
    def __init__(self, paragraph):
        self.paragraph = paragraph

    def p_replace(self, x:int, key:str, value:str, textimage=0, w=1, h=1):
        '''
        키워드에 포맷(볼드, 색깔 등)까지 반영하기 위해 직접  replace  하지 않고 run 활용
        :param x:           paragraph id
        :param key:         Keywords that need to be replaced
        :param value:       The replaced keywords
        :param textimage:   text replace or image insert, 0-text, 1-image
        :w : 이미지 일 경우 폭
        :h : 이미지 일 경우 높이
        :return:
        -하나의 문장에 있는 모든 문자의 위치를 구함 [{run_index , char_index}]
        -여기서 run이라 함은 하나의 문장에서 스타일이 달라지는 단위를 의미함. 예를들면 "나는 학교에 한다". 학교에가 볼드체이면 run은 "나는", "학교에", "간다"가 run 리스트
        -그런데 실제 한글과 영문,숫자 등이 있어 word 자체에서 결정한다. 즉 "학", "교에" 이렇게 될 수도 있다.
        '''
        p_maps = [{"run": y, "char": z} for y, run in enumerate(self.paragraph.runs) for z, char in enumerate(list(run.text))]

        # 아래는 찾고자 하는 key값의 한문장에서 첫번째 시작 인덱스 값 찾는다. k_idx=시작인덱스, 한문장에 여러개의 키워드 매칭이 있을 수 있으니 리스트 함
        k_idx = [s for s in range(len(self.paragraph.text)) if self.paragraph.text.find(key, s, len(self.paragraph.text)) == s]

        for i, start_idx in enumerate(reversed(k_idx)):       # 뒤에서 부터 해야지 인덱스가 변하지 않는다. 만약에 앞에서 부터 해버리면 뒤에 인덱스 달라진다. 키워드가 2개 이상 있을 때
            end_idx = start_idx + len(key)
            k_maps = p_maps[start_idx:end_idx]
            self.r_replace(k_maps, value, textimage)

    def r_replace(self, k_maps:list, value:str, textimage=0, w=1, h=1):
        '''
        여기 내용은 이런 것. 하나의 문자에서 문자 하나하나 지워야 스타일이 유지 된다. replace로 해버리면 기존 스타일 무시하고 디폴트 스타일로 문자를 대체한다.
        우리가 한글이든 word 든 문자에서 글자를 하나 지우고 다른 글자를 써야 그 스타일이 유지 된다. 단어자체를 바꾸기로 해버리면 볼드체 빨간색 등을 무시하고 디폴트로 되지 않는가
        '''
        for i, position in enumerate(reversed(k_maps), start=1):
            y, z = position["run"], position["char"]
            run:object = self.paragraph.runs[y]
            thisrun = list(run.text)
            if i < len(k_maps):
                thisrun.pop(z)
            if i == len(k_maps) and textimage==0:
                thisrun[z] = value
            if i == len(k_maps) and textimage==1: # 이미지 일경우
                run.text=""
                run.add_picture(value, width= Cm(w), height= Cm(h))
                break
            run.text = ''.join(thisrun)


# 단지 .docx만 지원됨. doc는 테스트 하지 않음
class WordReplace:

    def __init__(self, file):
        self.docx = Document(file)

    # 본문의 키워드 변경
    def body_content(self, replace_dict:dict,textimage=0,w=1, h=1):
        for key, value in replace_dict.items():
            for x, paragraph in enumerate(self.docx.paragraphs):
                Execute(paragraph).p_replace(x, key, value,textimage,w,h)
                
    # 본문 테이블, 자동 추가 없이 그냥 정해진 키워드 변경
    def body_tables(self,replace_dict:dict,textimage=0,w=1, h=1):
        for key, value in replace_dict.items():
            for table in self.docx.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for x, paragraph in enumerate(cell.paragraphs):
                            Execute(paragraph).p_replace(x, key, value,textimage,w,h)

    # 리스트 요소에 딕셔너리가 있을 경우. 이런 경우도 이미 테이블 로우가 이미 정의가 되어 있어야 함
    # 리스트 형식으로 자동 추가되게 하려면 엑셀로 하는게 낫겠지. 여기는 대장만 출력하는 용도
    def body_tables_list(self,replace_list:list,textimage=0,w=1, h=1):
        for value in replace_list:
            self.body_tables(value,textimage,w,h)

    def header_content(self,replace_dict:dict,textimage=0,w=1, h=1):
        for key, value in replace_dict.items():
            for section in self.docx.sections:
                for x, paragraph in enumerate(section.header.paragraphs):
                    Execute(paragraph).p_replace(x, key, value,textimage,w,h)


    def header_tables(self,replace_dict:dict,textimage=0,w=1, h=1):
        for key, value in replace_dict.items():
            for section in self.docx.sections:
                for table in section.header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for x, paragraph in enumerate(cell.paragraphs):
                                Execute(paragraph).p_replace(x, key, value,textimage,w,h)


    def footer_content(self, replace_dict:dict,textimage=0,w=1, h=1):
        for key, value in replace_dict.items():
            for section in self.docx.sections:
                for x, paragraph in enumerate(section.footer.paragraphs):
                    Execute(paragraph).p_replace(x, key, value,textimage,w,h)


    def footer_tables(self, replace_dict:dict,textimage=0,w=1, h=1):
        for key, value in replace_dict.items():
            for section in self.docx.sections:
                for table in section.footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for x, paragraph in enumerate(cell.paragraphs):
                                Execute(paragraph).p_replace(x, key, value,w,h)


    def save(self, filepath:str):
        self.docx.save(filepath)

    @staticmethod
    def docx_list(dirPath):
        fileList = []
        for roots, dirs, files in os.walk(dirPath):
            for file in files:
                if file.endswith("docx") and file[0] != "~":  # Find the docx document and exclude temporary files
                    fileRoot = os.path.join(roots, file)
                    fileList.append(fileRoot)
        return fileList