from typing import Iterable
from docx import Document
from docx.package import ImageParts
import openpyxl as op
import cx_Oracle
import datetime
import os, time
import platform
import SU_File_Convert
import Chemp_Bad_SQL as cbs

######################################################################## 데이터베이스 연결
# 기술원 개발서버
os.putenv('NLS_LANG', '.UTF8')
connection = cx_Oracle.connect('chemp/chemp@192.168.50.60:1521/orcl')
cursor = connection.cursor()

# 기술원 운영서버
# os.putenv('NLS_LANG', '.UTF8')
# connection = cx_Oracle.connect('chemp/chemp!1299@10.10.20.10:1521/chemp')
# cursor = connection.cursor()

# 로컬설버
# os.putenv('NLS_LANG', '.UTF8')
# connection = cx_Oracle.connect('scott/tiger@127.0.0.1:1521/XE')
# cursor = connection.cursor()


######################################################################## 기본정보 세팅
G_ExFilePos = os.getcwd()				# 현재실행위치
now = datetime.datetime.now()			# 날짜
strTime = now.strftime('%Y-%m-%d')		
G_Platform = platform.system()			# 플랫폼
G_SplitDef = ''							# 플랫폼 분리자

if G_Platform == 'Windows':
	G_ExFilePosIn = G_ExFilePos + '\\GIDT\\'
	G_ExFilePosOut = G_ExFilePos + '\\BADResult\\'
	G_SplitDef = '\\'
else:
	G_ExFilePosIn = G_ExFilePos + '/GIDT/'
	G_ExFilePosOut = G_ExFilePos + '/BADResult/'
	G_SplitDef = '/'
	
# Chemp_Bad_SQL에 존재하는 글로벌 함수 및 변수 초기화
cbs.G_BAD_CURSOUR 		= connection.cursor()
cbs.G_Local_Jpg   		= G_ExFilePos + G_SplitDef + "ImgDown" + G_SplitDef

######################################################################## 문서가져오기
rowcount = 0
while True:
	rowcount += 1

	# 여기는 모니터링 할 테이블, 특정이벤트 일때 모니터링 테이블에 삽입
	cursor.execute("""SELECT NUMID FROM TN_BADREPORT_MONITORING""")
	monrows = cursor.fetchall()

	# 출력버튼이 눌러졌을 때 액션 시작
	for imonrow in range(0, len(monrows)):
		keyvalue    = monrows[imonrow][0]	# 키값
		gtabcount 	= 1						# 시험자료 개수만큼 돌기
		fileList 	= []
		gmasterfile = ""

		# 시험항목마다 시험자료가 여러개인 경우가 있다. 메인은 가져오고 이후 서브에다 기록해서 문서를 하나로 합친다.
		# 출력해야 할 항목을 가져온다.
		BAD_ITEM = cbs.G_BAD_REPORTING_LIST
		for badsubitem in range(0,len(BAD_ITEM)):

			# 시험자료 1
			gmasterfile = BAD_ITEM[badsubitem][0] + ".docx"
			gbadoc = Document(G_ExFilePosIn + gmasterfile)
			gtable = gbadoc.tables  		# 테이블을 리스트 형식으로
			gpar   = gbadoc.paragraphs  	# 나머지 문장도 리스트 형식으로

			ibaddoc = gbadoc
			itable  = gtable
			ipar    = gpar

			# 시험자료가 여러개 일 경우 2부터 여기에 담는다.
			if BAD_ITEM[badsubitem][3] == 1:		# 서브항목이 있는다는 것
				gtabcount = 100
			else:
				gtabcount = 1						# 서브항목이 없음

			# 여기서 하나의 항목을 처리한다.
			for tab in range(0, gtabcount):

				# 값을 가져오는 테이블
				sql = cbs.Bad_Return_Sql(keyvalue, BAD_ITEM[badsubitem][1], BAD_ITEM[badsubitem][2], tab+1)
				cursor.execute(sql)
				rows = cursor.fetchall()

				if len(rows) > 0:

					# 정해진 코드에 값을 대체한다. 일반적으로 이렇게 사용할 듯
					# for pc in ipar:
					# 	pc.text = pc.text.replace("{01}", "테스트 문서")

					# 여기에서 문서에 적는다.
					cbs.Bad_Return_Rows(rows,BAD_ITEM[badsubitem][2],itable)

					# 시험자료 1은 마스터 파일
					if tab == 0:
						gmasterfile = G_ExFilePosOut + gmasterfile
						ibaddoc.save(gmasterfile)

						# 마스터 파일에 한번 기재하고 나면 이제부터는 서브파일에 기재하기 위해 바꿔준다. 1이면 서브파일이 존재
						if BAD_ITEM[badsubitem][3] == 1:
							subbadoc = Document(G_ExFilePosIn + BAD_ITEM[badsubitem][0] + "_1.docx")
							subtable = subbadoc.tables
							subpar   = subbadoc.paragraphs

					# 시험자료가 2개 이상 존재하면 2부터는 서브파일, 서브파일명은 마스트파일명_1.docx 이미 서식에 존재해야 함
					else:
						subfile = G_ExFilePosOut + BAD_ITEM[badsubitem][0] + "_" + str(tab + 1) + ".docx"
						ibaddoc.save(subfile)
						fileList.append(subfile)
				else:	# 더이상 시험자료 추가 없거나 값이 없을 경우
					break

			# 마스터 파일과 서브 파일이 존재하는 경우 서브 파일등을 합친다.
			if len(gmasterfile) > 0 and len(fileList) > 0:
				SU_File_Convert.sfc_combine_all_docx(gmasterfile, fileList, gmasterfile)

				# 리눅스에서는 먹지 않음
				l_pdffile = G_ExFilePosOut + BAD_ITEM[badsubitem][0] + ".pdf"
				r_pdffile = cbs.G_Remote_PdfResult + BAD_ITEM[badsubitem][0] + ".pdf"
				SU_File_Convert.sfc_wordtopdf(gmasterfile, l_pdffile)

				# 원격 운영서버로 최종결과물 전송
				cbs.Remote_FileUp(l_pdffile,r_pdffile)

				# 중간 작업파일 삭제
				cbs.Remote_FileRemove(G_ExFilePosOut)
			else:
				# 리눅스에서는 먹지 않음
				l_pdffile = G_ExFilePosOut + BAD_ITEM[badsubitem][0] + ".pdf"
				r_pdffile = cbs.G_Remote_PdfResult + BAD_ITEM[badsubitem][0] + ".pdf"
				SU_File_Convert.sfc_wordtopdf(gmasterfile, l_pdffile)

				# 원격 운영서버로 최종결과물 전송
				cbs.Remote_FileUp(l_pdffile, r_pdffile)

				# 중간 작업파일 삭제
				cbs.Remote_FileRemove(G_ExFilePosOut)

			# 문서가 완료되면 모니터링 테이블에서 삭제한다.
			cursor.execute("""DELETE FROM TN_BADREPORT_MONITORING WHERE NUMID=%s""" %(keyvalue))
			connection.commit()

			print("=============>" + str(keyvalue))

	print("message wait...")
	time.sleep(3)