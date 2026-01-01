from typing import Iterable
from docx import Document
from docx.package import ImageParts
import openpyxl as op
import cx_Oracle
import datetime
import os
import time
import SU_File_Convert
import Chemp_Bad_SQL_2022 as cbs


# 데이터베이스 연결
# 기술원 개발서버
# os.putenv('NLS_LANG', '.UTF8')
# connection = cx_Oracle.connect('chemp/chemp@192.168.50.60:1521/orcl')
# cursor = connection.cursor()

# 기술원 운영서버
# os.putenv('NLS_LANG', '.UTF8')
# connection = cx_Oracle.connect('chemp/chemp!1299@10.10.20.10:1521/chemp')
# cursor = connection.cursor()

# 로컬설버
os.putenv('NLS_LANG', '.UTF8')
connection = cx_Oracle.connect('scott/tiger@192.168.56.1:1521/XE')
cursor = connection.cursor()


# 기본정보 세팅
G_ExFilePos = os.getcwd()
now = datetime.datetime.now()
strTime = now.strftime('%Y-%m-%d')
G_SplitDef = ''

G_ExFilePosIn = G_ExFilePos + '\\BAD\\'
G_ExFilePosOut = G_ExFilePos + '\\BADResult\\'
G_SplitDef = '\\'

# Chemp_Bad_SQL에 존재하는 글로벌 함수 및 변수 초기화
cbs.G_BAD_CURSOUR = connection.cursor()
cbs.G_Local_Jpg = G_ExFilePos + G_SplitDef + "ImgDown" + G_SplitDef


# 보고서 메인 항목과 메인항목의 세부항목중 보고서에 존재하는 세부항목을 정의한다.
# BAD_ITEM은 메인항목, BAD_SUBITEM은 보고서에 존재하는 세부항목
wb = op.load_workbook(G_ExFilePosIn + "BAD.xlsx")
ws1 = wb.worksheets[0]  # 메인항목
ws2 = wb.worksheets[1]  # 서브항목
BAD_ITEM = []
BAD_SUBITEM = {}

for rowitem in range(2, ws1.max_row+1):
    listbad = []    # 메인항목
    for colitem in range(1, ws1.max_column+1):
        strVal = str(ws1.cell(row=rowitem, column=colitem).value).strip()
        listbad.append(strVal)

        # 메인항목일때만 서브항목을 구함
        if colitem == 1:
            dicbad = {}
            for srowitem in range(2, ws2.max_row+1):
                strVal2 = str(ws2.cell(row=srowitem, column=1).value).strip()
                strValKey = str(ws2.cell(row=srowitem, column=2).value).strip()
                strValFormat = str(
                    ws2.cell(row=srowitem, column=3).value).strip()
                if strVal == strVal2:
                    dicbad[strValKey] = strValFormat
            BAD_SUBITEM[strVal] = dicbad
    BAD_ITEM.append(listbad)
wb.close()


rowcount = 0
while True:
    rowcount += 1

    # 모니터링 테이블. 즉 사용자가 신청보고서 출력 요청하면 이 테이블에 값을 세팅
    cursor.execute("""SELECT NUMID FROM TN_BADREPORT_MONITORING""")
    monrows = cursor.fetchall()

    # 보고서 모니터링 요청을 한 숫자만큼 반복해서 보고서를 만들어 냄
    for imonrow in range(0, len(monrows)):
        keyvalue = monrows[imonrow][0]  # 키값
        gtabcount = 1						   # 시험자료 개수만큼 돌기. 일단 1부터 시작함
        fileList = []
        gmasterfile = ""

        for badsubitem in range(0, len(BAD_ITEM)):

            # 문서의 기본 정보를 세팅(문서명, 메인항목아이디, 서브항목아이디, 시험자료 첨부여부)
            baditem1 = BAD_ITEM[badsubitem][0]
            baditem2 = BAD_ITEM[badsubitem][1]
            baditem3 = BAD_ITEM[badsubitem][2]
            baditem4 = BAD_ITEM[badsubitem][3]

            # 메인항목명을 파일명으로 사용
            gmasterfile = baditem1 + ".docx"
            gbadoc = Document(G_ExFilePosIn + gmasterfile)
            gtable = gbadoc.tables
            gpar = gbadoc.paragraphs

            ibaddoc = gbadoc
            itable = gtable
            ipar = gpar

            # 항목에 시험자료가 있는 구조이면 기본적으로 시험자료 체크를 20번으로 함(1-서브시험자료 존재, 0-없음)
            # 사실 시험자료를 20개까지 추가하는 항목은 없음
            if baditem4 == 1:
                gtabcount = 20
            else:
                gtabcount = 1

            # 시험자료 개수만큼 세팅한다.
            for tab in range(0, gtabcount):

                # 메인항목 출력을 위한 쿼리를 메인항목.txt로 된 파일을 불러와서 조건을 replace 한다.
                # 메인항목마다 쿼리가 다 다르다고 한다. 때문에 파일로 하나씩 정리해서 불러오기 하는게 낫다.
                strList = [keyvalue,  baditem2, baditem3, tab+1]  # 바꿀문자열 리스트
                f = open(G_ExFilePosIn+baditem1+".txt", 'r')
                sql = f.read()
                for i in range(0, len(strList)):
                    sql = sql.replace("%s"+str(i+1), strList[i])
                f.close()

                cursor.execute(sql)
                rows = cursor.fetchall()

                if len(rows) > 0:

                    # 보고서에 문서를 적는다.

                    # 일반적인 값을 대처하는 방식(문장)
                    # for pc in ipar:
                    # 	pc.text = pc.text.replace("{01}", "테스트 문서")

                    # 보고서에 담을 메인항목에 해당하는 세부항목과 타입을 가져오기 위해
                    currDic = BAD_SUBITEM.get(baditem2)
                    
                    # 테이블형태에 값 쓰기
                    cbs.Bad_Return_Rows(rows, baditem3, currDic, itable)

                    # 시험자료 1은 마스터 파일
                    if tab == 0:
                        gmasterfile = G_ExFilePosOut + gmasterfile
                        ibaddoc.save(gmasterfile)

                        # 마스터 파일에 한번 기재하고 나면 이제부터는 서브파일에 기재하기 위해 바꿔준다. 1이면 서브파일이 존재
                        # 즉 아래는 무조건 한번만 실행. 아래와 같이 ibaddoc, itable, ipar 서브 파일로 할당
                        if BAD_ITEM[badsubitem][3] == 1:
                            ibaddoc = Document(G_ExFilePosIn + baditem1 + "_1.docx")
                            itable = ibaddoc.tables
                            ipar = ibaddoc.paragraphs

                    # 시험자료가 2개 이상 존재하면 2부터는 서브파일, 서브파일명은 마스트파일명_1.docx 이미 서식에 존재해야 함
                    else:
                        subfile = G_ExFilePosOut + baditem1 + "_" + str(tab + 1) + ".docx"
                        ibaddoc.save(subfile)
                        fileList.append(subfile)

                else:
                    break

            # 마스터 파일과 서브 파일이 존재하는 경우 서브 파일등을 합친다.
            if len(gmasterfile) > 0 and len(fileList) > 0:
                SU_File_Convert.sfc_combine_all_docx(
                    gmasterfile, fileList, gmasterfile)

                # 워드 파일을 PDF로 변환 리눅스에서는 먹지 않음
                l_pdffile = G_ExFilePosOut + BAD_ITEM[badsubitem][0] + ".pdf"
                r_pdffile = cbs.G_Remote_PdfResult + BAD_ITEM[badsubitem][0] + ".pdf"
                SU_File_Convert.sfc_wordtopdf(gmasterfile, l_pdffile)

                # 원격 운영서버로 최종결과물 전송
                cbs.Remote_FileUp(l_pdffile, r_pdffile)

                # 중간 작업파일 삭제
                cbs.Remote_FileRemove(G_ExFilePosOut)
            else:
                # 리눅스에서는 먹지 않음
                l_pdffile = G_ExFilePosOut + BAD_ITEM[badsubitem][0] + ".pdf"
                r_pdffile = cbs.G_Remote_PdfResult + BAD_ITEM[badsubitem][0] + ".pdf"
                SU_File_Convert.sfc_wordtopdf(gmasterfile, l_pdffile)

                # 원격 운영서버로 최종결과물 전송
                cbs.Remote_FileUp(l_pdffile, r_pdffile)

                # 중간 작업파일 삭제
                cbs.Remote_FileRemove(G_ExFilePosOut)

            # 문서가 완료되면 모니터링 테이블에서 삭제한다.
            cursor.execute("""DELETE FROM TN_BADREPORT_MONITORING WHERE NUMID=%s""" % (keyvalue))
            connection.commit()

            print("=============>" + str(keyvalue))

    print("message wait...")
    time.sleep(3)
