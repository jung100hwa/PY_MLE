from typing import Iterable
from docx import Document
from docx.package import ImageParts
import openpyxl as op
import cx_Oracle
import datetime
import os
import time
import SU_File_Convert
import Chemp_Bad_SQL_2023 as cbs




# 데이터베이스 연결
# 기술원 개발서버
# os.putenv('NLS_LANG', '.UTF8')
# connection = cx_Oracle.connect('chemp/chemp@192.168.50.60:1521/orcl')
# cursor = connection.cursor()

# 기술원 운영서버
os.putenv('NLS_LANG', '.UTF8')
connection = cx_Oracle.connect('chemp/chemp!1299@10.10.20.10:1521/chemp')
cursor = connection.cursor()

# 로컬설버
# os.putenv('NLS_LANG', '.UTF8')
# connection = cx_Oracle.connect('scott/tiger@192.168.56.1:1521/XE')
# cursor = connection.cursor()


# 기본정보 세팅
G_ExFilePos = os.getcwd()
now          = datetime.datetime.now()
strTime      = now.strftime('%Y-%m-%d')

G_ExFilePosIn = G_ExFilePos + '\\BAD\\'
G_ExFilePosOut = G_ExFilePos + '\\BADResult\\'

# Chemp_Bad_SQL에 존재하는 글로벌 함수 및 변수 초기화
cbs.G_BAD_CURSOUR = connection.cursor()
cbs.G_Local_Jpg = G_ExFilePos + "\\ImgDown\\"


# 보고서 메인 항목과 메인항목의 세부항목중 보고서에 존재하는 세부항목을 정의한다.
# BAD_ITEM은 메인항목, 세부항목은 판다스
wb = op.load_workbook(G_ExFilePosIn + "BAD.xlsx")
ws1 = wb.worksheets[0]  # 메인항목
BAD_ITEM = []

for rowitem in range(2, ws1.max_row+1):
    listbad = []
    for colitem in range(1, ws1.max_column+1):
        strVal = str(ws1.cell(row=rowitem, column=colitem).value).strip()
        listbad.append(strVal)
    BAD_ITEM.append(listbad)
wb.close()


rowcount = 0
while True:
    rowcount += 1

    # 모니터링 테이블. 즉 사용자가 신청보고서 출력 요청하면 이 테이블에 값을 세팅
    # cursor.execute("""SELECT NUMID FROM TN_BADREPORT_MONITORING""")
    cursor.execute("""SELECT CAS_NO FROM HGL WHERE CAS_NO='994'""")
    monrows = cursor.fetchall()

    # 보고서 모니터링 요청을 한 숫자만큼 반복해서 보고서를 만들어 냄
    for imonrow in range(0, len(monrows)):
        keyvalue = monrows[imonrow][0]  # 키값
        gtabcount = 1						   # 시험자료 개수만큼 돌기. 일단 1부터 시작함
        fileList = []
        gmasterfile = ""

        for bitem in range(0, len(BAD_ITEM)):

            baditem1 = BAD_ITEM[bitem][0]  # 문서명
            baditem2 = BAD_ITEM[bitem][1]  # 메인아이디
            baditem3 = BAD_ITEM[bitem][2]  # 서브아이디
            baditem4 = BAD_ITEM[bitem][3]  # 시험자료 첨부여부

            # 메인항목명을 파일명으로 사용
            gmasterfile = baditem1 + ".docx"
            gbadoc = Document(G_ExFilePosIn + gmasterfile)
            gtable = gbadoc.tables
            gpar = gbadoc.paragraphs

            ibaddoc = gbadoc
            itable = gtable
            ipar = gpar

            # 항목에 시험자료가 있는 구조이면 기본적으로 시험자료 체크를 10번으로 함(1-서브시험자료 존재, 0-없음)
            # 사실 시험자료를 20개까지 추가하는 항목은 없음
            if baditem4 == 1:
                gtabcount = 10
            else:
                gtabcount = 1

            # 시험자료 개수만큼 세팅한다.
            for tab in range(0, gtabcount):

                # 메인항목 출력을 위한 쿼리를 메인항목.txt로 된 파일을 불러와서 조건을 replace 한다.
                # 메인항목마다 쿼리가 다 다르다고 한다. 때문에 파일로 하나씩 정리해서 불러오기 하는게 낫다.
                # 메인아이디, 키(물질승인신청의 no), 시험자료 추가할때 인덱스, 서브아이디
                strList = [baditem2, keyvalue, tab+1, baditem3]  # 바꿀문자열 리스트
                f = open(G_ExFilePosIn+"BAD_01.txt", 'r')
                sql = f.read()
                
                for i in range(0, len(strList)):
                    sql = sql.replace("%s"+str(i+1), str(strList[i]))
                f.close()

                cursor.execute(sql)
                rows = cursor.fetchall()

                if len(rows) > 0:

                    # 테이블형태에 값 쓰기
                    cbs.Bad_Return_Rows(rows, itable, ipar, baditem2)

                    # 시험자료 1은 마스터 파일
                    if tab == 0:
                        gmasterfile = G_ExFilePosOut + gmasterfile
                        ibaddoc.save(gmasterfile)

                        # 마스터 파일에 한번 기재하고 나면 이제부터는 서브파일에 기재하기 위해 바꿔준다. 1이면 서브파일이 존재
                        # 즉 아래는 무조건 한번만 실행. 아래와 같이 ibaddoc, itable, ipar 서브 파일로 할당
                        if baditem4 == 1:
                            ibaddoc = Document(G_ExFilePosIn + baditem1 + "_1.docx")
                            itable = ibaddoc.tables
                            ipar = ibaddoc.paragraphs

                    # 시험자료가 2개 이상 존재하면 2부터는 서브파일, 서브파일명은 마스트파일명_1.docx 이미 서식에 존재해야 함
                    else:
                        subfile = G_ExFilePosOut + baditem1 + "_" + str(tab + 1) + ".docx"
                        ibaddoc.save(subfile)
                        fileList.append(subfile)

                else:
                    break

            # 마스터 파일과 서브 파일이 존재하는 경우 서브 파일등을 합친다.
            if len(gmasterfile) > 0 and len(fileList) > 0:
                SU_File_Convert.sfc_combine_all_docx(
                    gmasterfile, fileList, gmasterfile)

                # 워드 파일을 PDF로 변환 리눅스에서는 먹지 않음
                l_pdffile = G_ExFilePosOut + baditem1 + ".pdf"
                r_pdffile = cbs.G_Remote_PdfResult + baditem1 + ".pdf"
                SU_File_Convert.sfc_wordtopdf(gmasterfile, l_pdffile)

                # # 원격 운영서버로 최종결과물 전송
                # cbs.Remote_FileUp(l_pdffile, r_pdffile)
                #
                # # 중간 작업파일 삭제
                # cbs.Remote_FileRemove(G_ExFilePosOut)
            else:
                # 리눅스에서는 먹지 않음
                l_pdffile = G_ExFilePosOut + baditem1 + ".pdf"
                r_pdffile = cbs.G_Remote_PdfResult + baditem1 + ".pdf"
                SU_File_Convert.sfc_wordtopdf(gmasterfile, l_pdffile)

                # # 원격 운영서버로 최종결과물 전송
                # cbs.Remote_FileUp(l_pdffile, r_pdffile)
                #
                # # 중간 작업파일 삭제
                # cbs.Remote_FileRemove(G_ExFilePosOut)

        # 문서가 완료되면 모니터링 테이블에서 삭제한다.
        # cursor.execute("""DELETE FROM TN_BADREPORT_MONITORING WHERE NUMID=%s""" % (keyvalue))
        cursor.execute("""DELETE FROM HGL WHERE CAS_NO='%s'""" % (keyvalue))
        connection.commit()

        print("=============>" + str(keyvalue))

    print("message wait...")
    time.sleep(3)
