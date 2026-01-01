import os
import paramiko
from stat import S_ISDIR
import datetime
import glob
from docx.shared import Inches
import pandas as pd

# 세부 컬럼 등 정의된 엑셀을 판다스에 담는다.
G_Loc = os.getcwd()
G_Loc = G_Loc + '\\BAD\\SUBBAD.xlsx'
subdf = pd.read_excel(G_Loc)

# 운영서버 환경정의
Remote_Server = "210.220.13.165"
Remote_UserName = "chkusr"
Remote_UserPwd = "Qwer!234"
Remote_port = 2022
Remote_Jpg = '/CPMS/WAS/reg/files/'
Local_Jpg  = "c:\\work\\"

# 개발서버 환경정의
# Remote_Server = "192.168.50.60"
# Remote_UserName = "chkusr"
# Remote_UserPwd = "Qwer!234"
# Remote_port = 2022
# Remote_Jpg = '/CPMS/WAS/reg/files/'

global G_Remote_PdfResult		        # pdf 최종보고서 저장
global G_BAD_CURSOUR
global G_Local_Jpg

# 원격서버 리포팅 결과물을 저장할 결과물 폴더
G_Remote_PdfResult = '/CPMS/WAS/reg/reporttest/'

# 구조식에 이미지 정보를 다운로드 하기 위한 것
ssh = paramiko.SSHClient()
ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
ssh.connect(Remote_Server, username=Remote_UserName,
            password=Remote_UserPwd, port=Remote_port)
sftp = ssh.open_sftp()


# 문서의 테이블 안에 값을 세팅
# rows         : 쿼리 결과물
# itable        : 보고서 내의 테이블
# ipar          : 문장전체(테이블 아닐 경우 replace 하기 위해)
# mainkey    : 메인항목 키
def Bad_Return_Rows(rows, itable, ipar, mainkey):

    count = 0
    rowcount = 0
    
    # 메인키에 해당되는 것만 불러온다.    
    mask = "maincode==@mainkey"
    ndf = subdf.query(mask)

    for itemrow in rows:

        # 쿼리 결과 항목 중 보고서에 출력할 항목인지 일단 검토
        # returnVal 은 엑셀에 정의한 컬럼이 코드인지, 그림인지 등을 불러온다.
        strSubcode = str(itemrow[1]).strip()
        strColnm = str(itemrow[2]).strip()
        
        smask = "(subcode==@strSubcode) and (colname==@strColnm)"
        idf = ndf.query(smask)
        returnVal = idf.iloc[0,4]       # valueformat

        if returnVal:

            if returnVal == '1':      # 텍스트값
                itable[0].rows[count].cells[1].text = str(itemrow[17])

            if returnVal == '2':      # 코드값
                codevalue = Bad_Return_SqlkeyvalueBAD(
                    str(itemrow[11]), str(itemrow[17]))
                itable[0].rows[count].cells[1].text = codevalue

            if returnVal == '3':      # 범위값
                vlist = Bad_Return_Split(str(itemrow[17]), "##")
                itable[0].rows[count].cells[2].text = vlist[0]
                itable[0].rows[count].cells[4].text = vlist[1]

            if returnVal == '4':      # 이미지 첨부파일(구조식 같은)
                codevalue = Remote_FileDown(str(itemrow[27]), str(itemrow[27]))
                itable[0].rows[count].cells[1]._element.clear_content()
                img = itable[0].rows[count].cells[1].add_paragraph()
                r = img.add_run()
                r.add_picture(codevalue, width=Inches(2.5))

            # 첨부파일 구간반복 시작
            if returnVal == '5':
                imsicount = rowcount + 2
                gugan = ''
                while True:
                    strfilename = str(rows[imsicount][28]).strip()
                    strVal = str(rows[imsicount][7]).strip()

                    if strVal == '구간반복 끝':
                        break

                    gugan = gugan + strfilename + ", "
                    imsicount = imsicount + 1

                if len(gugan) > 2:
                    gugan = gugan.strip()
                    gugan = gugan[:-1]
                for pc in ipar:
                    pc.text = pc.text.replace("{01}",gugan)

            # 테이블에 일반적인 반복 형태
            if returnVal[0:1] == '6' and gubunnum == 'RG3171':
                imsicount          = rowcount + 2
                tabid              = int(returnVal[1:3])-1   # 테이블아이디
                tabcolallcount     = int(returnVal[3:5])     # 컬럼개수
                tabcolindex    = 1
                tabrowindxe    = 1
                while True:
                    strVal = str(rows[imsicount][7]).strip()
                    strVal2 = str(rows[imsicount][17]).strip()

                    if strVal2.lower() == "n/a" :
                        strVal2 = strVal2.replace('/','')

                    if strVal == '구간반복 끝':
                        break

                    itable[tabid].rows[tabrowindxe].cells[tabcolindex].text = strVal2

                    imsicount = imsicount + 1
                    tabcolindex = tabcolindex + 1

                    # 다음 로우이면 cell 인덱스를 초기화 한다. 다시 1부터 시작한다.
                    if tabcolallcount == tabcolindex:
                        tabrowindxe = tabrowindxe + 1
                        tabrowcount = 1

            count = count + 1
        rowcount = rowcount + 1
    return


# 순도처러럼 범위로 입력하는 것에 대한 값
# value      : splitvalue로 분리된 값, 예를들면 순도는 90##100
# splitvalue : 구분자, 순도와 같은 범위는 일반적으로 '##' 으로 구분되어 있음
def Bad_Return_Split(value, splitvaule):
    onevalue = ''
    twovaue = ''
    if len(str(value)) > 0 and len(str(splitvaule)) > 0:
        valueList = str(value).split(splitvaule)
        if len(str(valueList[0])) > 0:
            onevalue = valueList[0]
        else:
            onevalue = ""

        if len(str(valueList[1])) > 0:
            twovalue = valueList[1]
        else:
            twovalue = ""
    return onevalue, twovalue


# 키값 구하기, 공통코드값을 구하는 것. BAD만을 위한 코드값을 따로 정의
# groupcode : 그룹코드(상단에 정의되어 있음)
# itemcode   : 상세코드
def Bad_Return_SqlkeyvalueBAD(groupcode, itemcode):
    sql = """SELECT KOREAN_EXMP_NM FROM TC_RP_OECD_PICK_CODE_I WHERE EXMP_GROUP_CODE='%s' AND EXMP_CODE = '%s'""" % (
        groupcode, itemcode)
    G_BAD_CURSOUR.execute(sql)
    result = G_BAD_CURSOUR.fetchone()
    if result:
        return result[0]
    else:
        return ''

# 일반적인 코드값
# groupcode : 그룹코드(상단에 정의되어 있음)
# itemcode   : 상세코드
def Bad_Return_Sqlkeyvalue(groupcode, itemcode):
    sql = """SELECT GET_CODE_NM('%s','%s') FROM DUAL""" % (groupcode, itemcode)
    G_BAD_CURSOUR.execute(sql)
    result = G_BAD_CURSOUR.fetchone()
    if result:
        return result[0]
    else:
        return ''

# 주로 원격에 있는 이미지 파일을 다운로드 받아서 문서에 첨부한다.
def Remote_FileDown(R_filename, L_filename):
    # 임시 파일을 일단 지운다.(주로 구조식 등 임시파일). 많아지면 관리가 안됨
    try:
        R_filename = Remote_Jpg + R_filename
        L_filename = G_Local_Jpg + L_filename + ".jpg"
        sftp.get(R_filename, L_filename)
        return L_filename
    except:
        return ""


# 최종 파일(PDF) 운영서버 정해진 폴더에 업로드 한다.
# L_file : 로컬파일
# R_file : 운영서버 전송 파일
def Remote_FileUp(L_File, R_File):
    try:
        sftp.put(L_File, R_File)         # 서버쪽으로 전송
        os.remove(L_File)                # 로컬결과물 삭제
    except:
        return ""
    return ""


# 이미지 첨부를 위해 임시 다운로드파일 삭제.  파일만 삭제한다.
# 이 기능은 일단 실행을 보류한다. 동시작업이 있을 경우 동시성제어에 문제가 있을 수 있다.
# Target : 다운로드된 임시 파일(보통 이미지 파일)
def Remote_FileRemove(Target):
    try:
        for file in glob.iglob(Target + "**/*", recursive=True):
            if not os.path.isdir(file):
                if os.path.isfile(file):
                    os.remove(file)
    except:
        pass