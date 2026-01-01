import os
import paramiko
from stat import S_ISDIR
import datetime
import glob
from docx.shared import Inches

# 운영서버 환경정의
# Remote_Server = "210.220.13.165"
# Remote_UserName = "chkusr"
# Remote_UserPwd = "Qwer!234"
# Remote_port = 2022
# Remote_Jpg = '/CPMS/WAS/reg/files/'
# Local_Jpg  = "c:\\work\\"

# 개발서버 환경정의
Remote_Server = "192.168.50.60"
Remote_UserName = "chkusr"
Remote_UserPwd = "Qwer!234"
Remote_port = 2022
Remote_Jpg = '/CPMS/WAS/reg/files/'

global G_Remote_PdfResult		        # pdf 최종보고서 저장
global G_BAD_CURSOUR
global G_Local_Jpg

# 원격서버 리포팅 결과물을 저장할 결과물 폴더
G_Remote_PdfResult = '/CPMS/WAS/reg/reporttest/'

# 구조식에 이미지 정보를 다운로드 하기 위한 것
ssh = paramiko.SSHClient()
ssh.set_missing_host_key_policy(paramiko.AutoAddPolicy())
ssh.connect(Remote_Server, username=Remote_UserName,
            password=Remote_UserPwd, port=Remote_port)
sftp = ssh.open_sftp()


# 출력할 보고서를 정의(지속적으로 업데이트 필요)
# 첫번째 : 문서명
# 두번째 : 항목에 대한 고유번호
# 세번째 : 서브항목이 존재하는지 판단. 즉 시험자료 추가가 있는 항목인지(0:없음, 1:있음)
G_BAD_REPORTING_LIST = [
    ["BAD_2.1.1.1.살생물물질 식별정보", "RG1501", "RG3166-1", 0]
]

# 필요한 값만 가져오기
# 파일명_1 : 문서에 세팅할 데이터, 파일명_2 : 코드화 되어 있는 항목
# 2.1.1.1 살생물물질 식별정보
G_BAD_REPORTING_LIST_RG3166_1 = ['살생물물질명 (ISO 일반명)', '고유번호 (CAS No.)', '그 외 고유번호 (CAS No.)', '기존화학물질(KE) 번호',
                                 '유독물질 번호', '살생물제품 유형', '물질 식별 연관 형태', '순도 범위 (%(W/W))', '분자식', '분자량 (g/mol)',
                                 '몰 질량 (g/mol)', 'SMILES 표기법', '구조식', '광학 활성 및 이성질체에 대한 상세정보']
G_BAD_REPORTING_LIST_RG3166_2 = [['살생물제품 유형', '코드'], [
    '물질 식별 연관 형태', '코드'], ['구조식', '이미지'], ['순도 범위 (%(W/W))', '##']]


# keynum           : 승인신청서 넘버(no)
# tabnum           : 시험자료가 2개이상인 경우
# maingubunnum : 항목의 메인 고유넘버
# gubunnum       : 시험자료항목마다 고유넘버
def Bad_Return_Sql(keynum, maingubunum, gubunnum, tabnum):
    sql = """SELECT
	T2.NO,
	T1.ATRB_SN,
	T2.GROUP_INDEX,
	T2.SEQ,
	T1.EXPR_IEM_CODE,
	T1.EXPR_IEM_SE_CODE,
	T1.CLAS_NM,
	T1.ATRB_NM,
	T1.ENG_ATRB_NM,
	T1.FIELD_TYPE,
	T1.FIELD_SIZE,
	T1.EXPR_GROUP_CODE,
	T1.ORDR,
	T1.GROUP_ID,
	T1.FIELD_DATA_TYPE,
	T1.FIELD_DATA_SIZE,
	T1.SHOW_ETC_YN,
	T2.ATRB_VALUE,
	T1.REQUIRED,
	T1.ROWSPAN,
	T1.COLSPAN,
	T1.ROWSPAN_TITLE,
	T1.CHECK_COU,
	T1.DATA_PROTECT_YN,
	T1.DATA_SE_YN,
	T2.ATCH_FILE_ID,
	(
	SELECT
		F1.FILE_SN
	FROM
		COMTNFILEDETAIL F1
	WHERE
		T2.ATCH_FILE_ID = F1.ATCH_FILE_ID ) AS FILE_SN,
	(
	SELECT
		F1.STRE_FILE_NM
	FROM
		COMTNFILEDETAIL F1
	WHERE
		T2.ATCH_FILE_ID = F1.ATCH_FILE_ID ) AS STRE_FILE_NM,
	(
	SELECT
		F1.ORIGNL_FILE_NM
	FROM
		COMTNFILEDETAIL F1
	WHERE
		T2.ATCH_FILE_ID = F1.ATCH_FILE_ID ) AS FILE_INFO,
	NVL(T2.TAB_EXAM_INDEX_INFO, (SELECT EXAMNAME FROM TN_NBM_EXPR_IEM_I I1 WHERE I1.EXPR_IEM_CODE = T1.EXPR_IEM_CODE) ) AS TAB_EXAM_INDEX_INFO,
	T1.WINOPEN1,
	T1.WINOPEN2
FROM
	TN_NBM_EXPR_IEM_TMPLAT_ATRB_D T1,
	TN_NBM_EXPR_IEM_ATRB_D T2
WHERE
	T1.ATRB_SN = T2.ATRB_SN(+)
	AND T1.EXPR_IEM_CODE = T2.EXPR_IEM_CODE(+)
	AND T1.EXPR_IEM_SE_CODE = T2.EXPR_IEM_SE_CODE(+)
	AND T1.EXPR_IEM_CODE = '%s'
	AND %s = T2.NO(+)
	AND %s = T2.TAB_EXAM_INDEX (+)
	AND T1.EXPR_IEM_SE_CODE = '%s'
ORDER BY
	T2.TAB_EXAM_INDEX ASC,
	T2.SEQ,
	T1.ATRB_SN""" % (maingubunum, keynum, tabnum, gubunnum)
    return sql


# 문서의 테이블 안에 값을 세팅
# rows         : 쿼리 결과물
# gubunnum : 시험자료항목마다 고유넘버
def Bad_Return_Rows(rows, gubunnum, itable):
    count = 0
    for itemrow in rows:
        if gubunnum == "RG3166-1":
            # 정의된 항목이 존재하는지 확인하고
            returnvalue = Bad_Return_ListValue(
                itemrow[7], G_BAD_REPORTING_LIST_RG3166_1)
            if returnvalue:
                # 코드 형태인지  이미지 형태인지 검색해서 코드이면 코드값을 세팅한다.
                codevalue = Bad_Return_ListValuecode(itemrow[7], str(itemrow[17]), str(
                    itemrow[11]), str(itemrow[27]), G_BAD_REPORTING_LIST_RG3166_2)
                if codevalue[1] == "0":         # 코드값
                    itable[0].rows[count].cells[1].text = codevalue[0]
                elif codevalue[1] == "1":       # 이미지(구조식 같은)
                    itable[0].rows[count].cells[1]._element.clear_content()
                    img = itable[0].rows[count].cells[1].add_paragraph()
                    r = img.add_run()
                    r.add_picture(codevalue[0], width=Inches(2.5))
                elif codevalue[1] == "2":
                    vlist = Bad_Return_Split(str(itemrow[17]), codevalue[0])
                    itable[0].rows[count].cells[2].text = vlist[0]
                    itable[0].rows[count].cells[4].text = vlist[1]
                else:                           # 일반적 쿼리 결과값
                    itable[0].rows[count].cells[1].text = str(itemrow[17])
                count = count + 1
    return


# 순도처러럼 범위로 입력하는 것에 대한 값
# value : splitvalue로 분리된 값, 예를들면 순도는 90##100
# splitvalue : 구분자, 순도와 같은 범위는 일반적으로 '##' 으로 구분되어 있음
def Bad_Return_Split(value, splitvaule):
    onevalue = ''
    twovaue = ''
    if len(str(value)) > 0 and len(str(splitvaule)) > 0:
        valueList = str(value).split(splitvaule)
        if len(str(valueList[0])) > 0:
            onevalue = valueList[0]
        else:
            onevalue = ""

        if len(str(valueList[1])) > 0:
            twovalue = valueList[1]
        else:
            twovalue = ""
    return onevalue, twovalue

# 리스트 안에 해당 값이 존재하는지 검토, 즉 쿼리 값 중 보고서 출력에 필요한 항목 검사
# 결과값의 모든 항목을 보고서로 출력하지 않음


def Bad_Return_ListValue(value, badList):
    if len(value) > 0:
        for item in badList:
            value = str(value).replace(' ', '')
            item = str(item).replace(' ', '')
            if value == item:
                return value
    return 0


# 리스트 안에 있으면서 코드형인지 아니면 이미지형 인지 판단.
# value      : 코드값으로 표시할 라벨명
# vcode      : 코드값
# groupvcode : 코드값일 경우 그룹코드값
# imgname    : 구조식 처럼 이미지인 경우 실제 스토리지 저장파일 이름
# badlist    : 코드를 값으로 표시할 리스트(상단정의)
def Bad_Return_ListValuecode(value, vcode, groupvcode, imgname, badList):
    if len(value) > 0:
        for itemcount in range(0, len(badList)):
            value = str(value).replace(' ', '')
            item = str(badList[itemcount][0]).replace(
                ' ', '')                  # 항목명
            itemcodeorimage = str(badList[itemcount][1]).replace(
                ' ', '')                  # 이미지 일경우 "이미지"
            if (value == item) and itemcodeorimage == '코드':                              # 일반적인 코드를 네임으로 변경
                codevalue = Bad_Return_SqlkeyvalueBAD(groupvcode, vcode)
                if codevalue:
                    return codevalue, "0"
                else:
                    return "", ""
            elif (value == item) and itemcodeorimage == '이미지' and len(imgname) > 0:     # 구조식 같은 이미지 출력을 위한 것
                codevalue = Remote_FileDown(imgname, imgname)
                if len(codevalue) > 0:
                    return codevalue, "1"
                else:
                    return "", ""
            elif (value == item) and itemcodeorimage == '##':        # 순도와 같이 범위 입력
                return "##", "2"
    return "", ""

# 키값 구하기, 공통코드값을 구하는 것. BAD만을 위한 코드값을 따로 정의
# groupcode : 그룹코드(상단에 정의되어 있음)
# itemcode  : 상세코드


def Bad_Return_SqlkeyvalueBAD(groupcode, itemcode):
    sql = """SELECT KOREAN_EXMP_NM FROM TC_RP_OECD_PICK_CODE_I WHERE EXMP_GROUP_CODE='%s' AND EXMP_CODE = '%s'""" % (
        groupcode, itemcode)
    G_BAD_CURSOUR.execute(sql)
    result = G_BAD_CURSOUR.fetchone()
    if result:
        return result[0]
    else:
        return ''

# 일반적인 코드값
# groupcode : 그룹코드(상단에 정의되어 있음)
# itemcode  : 상세코드


def Bad_Return_Sqlkeyvalue(groupcode, itemcode):
    sql = """SELECT GET_CODE_NM('%s','%s') FROM DUAL""" % (groupcode, itemcode)
    G_BAD_CURSOUR.execute(sql)
    result = G_BAD_CURSOUR.fetchone()
    if result:
        return result[0]
    else:
        return ''

# 주로 원격에 있는 이미지 파일을 다운로드 받아서 문서에 첨부한다.


def Remote_FileDown(R_filename, L_filename):
    # 임시 파일을 일단 지운다.(주로 구조식 등 임시파일). 많아지면 관리가 안됨
    Remote_FileRemove(G_Local_Jpg)
    try:
        R_filename = Remote_Jpg + R_filename
        L_filename = G_Local_Jpg + L_filename + ".jpg"
        sftp.get(R_filename, L_filename)
        return L_filename
    except:
        return ""
    return ""


# 최종 파일(PDF) 운영서버 정해진 폴더에 업로드 한다.
# L_file : 로컬파일
# R_file : 운영서버 전송 파일
def Remote_FileUp(L_File, R_File):
    try:
        sftp.put(L_File, R_File)         # 서버쪽으로 전송
        os.remove(L_File)                # 로컬결과물 삭제
    except:
        return ""
    return ""


# 이미지 첨부를 위해 임시 다운로드파일 삭제.  파일만 삭제한다.
# Target : 다운로드된 임시 파일(보통 이미지 파일)
def Remote_FileRemove(Target):
    try:
        for file in glob.iglob(Target + "**/*", recursive=True):
            if not os.path.isdir(file):
                if os.path.isfile(file):
                    os.remove(file)
    except:
        pass
